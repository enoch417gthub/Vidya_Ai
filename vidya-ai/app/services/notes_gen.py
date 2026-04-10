# app/services/notes_gen.py
# ============================================================
# Generates structured study notes from the knowledge base.
# Uses RAG to retrieve chapter content, then prompts LLM to
# format it as different note types (summary, bullets, Q&A).
# Exports to .docx using python-docx.
# ============================================================
import os
from datetime import datetime
from typing import Literal
from loguru import logger
from docx import Document as WordDoc
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.retriever import hybrid_search
from app.core.llm_engine import generate_response

# Note type options available to the student
NoteType = Literal['summary', 'bullets', 'qa', 'concepts', 'exam_ready']


def build_notes_prompt(topic: str, context: str, note_type: NoteType) -> str:
    '''
    Build the LLM prompt for generating study notes.
    Different note_types produce different output formats.
    '''
    format_instructions = {
        'summary': (
            'Write a clear, 3-5 paragraph summary of the topic. '
            'Use simple language suitable for a student. '
            'Cover the main concepts in order.'
        ),
        'bullets': (
            'List ALL key points as short bullet points. '
            'Group related points under sub-headings. '
            'Each bullet must be one clear, complete idea.'
        ),
        'qa': (
            'Generate 10 important Questions and Answers from this content. '
            'Format strictly as: Q1: [question] A1: [answer] Q2: etc. '
            'Cover definitions, processes, and applications.'
        ),
        'concepts': (
            'List all key concepts, definitions, and formulas from this content. '
            'Format: TERM: definition. For formulas use: FORMULA: expression -- meaning. '
            'Sort alphabetically.'
        ),
        'exam_ready': (
            'Create an exam-ready cheat sheet. Include: '
            '1) Must-Know Definitions (5-10 terms), '
            '2) Important Formulas/Laws, '
            '3) Common Exam Question Patterns, '
            '4) Quick Revision Bullets (10 points). '
        ),
    }

    instruction = format_instructions.get(note_type, format_instructions['summary'])

    return f'''<|begin_of_text|><|start_header_id|>system<|end_header_id|>
You are an expert study notes creator. Create notes strictly from the provided context.
Do not add information not present in the context.<|eot_id|>
<|start_header_id|>user<|end_header_id|>
TOPIC: {topic}
CONTENT FROM TEXTBOOK:
{context}
TASK: {instruction}<|eot_id|>
<|start_header_id|>assistant<|end_header_id|>'''


def generate_notes(
    grade: str,
    subject: str,
    topic: str,
    note_type: NoteType = 'summary',
    language: str = 'en'
) -> dict:
    '''
    Generate study notes for a topic and save as a Word document.
    Returns: {content: str, export_path: str, word_doc_path: str}
    '''
    kb_dir = os.getenv('KNOWLEDGE_BASE_DIR', './knowledge_base')
    index_path = os.path.join(kb_dir, grade, subject, f'{subject}.faiss')

    if not os.path.exists(index_path):
        raise FileNotFoundError(f'No index for {grade}/{subject}')

    # Step 1: Retrieve relevant chunks for the topic
    logger.info(f'Generating {note_type} notes for: {topic}')
    chunks = hybrid_search(index_path, topic, top_k=8)  # More chunks for notes
    context = '\n\n'.join([c['content'] for c in chunks])

    if not context.strip():
        return {'content': 'No content found for this topic.', 'export_path': None}

    # Step 2: Generate notes via LLM
    prompt = build_notes_prompt(topic, context, note_type)
    notes_content = generate_response(prompt, max_tokens=800)

    # Step 3: Export to Word document
    export_path = export_notes_to_docx(
        topic=topic, subject=subject, grade=grade,
        content=notes_content, note_type=note_type
    )

    return {
        'content': notes_content,
        'export_path': export_path,
        'topic': topic,
        'note_type': note_type
    }


def export_notes_to_docx(
    topic: str, subject: str, grade: str,
    content: str, note_type: str
) -> str:
    '''
    Save generated notes as a formatted .docx Word document.
    Returns the file path of the saved document.
    '''
    doc = WordDoc()  # Create a new Word document

    # --- Title ---
    title = doc.add_heading(f'VIDYA AI Study Notes', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading(f'{topic.title()} -- {subject.title()} ({grade})', level=1)

    # --- Metadata line ---
    meta = doc.add_paragraph()
    meta.add_run(f'Note Type: {note_type.upper()} | Generated: {datetime.now().strftime("%d %b %Y, %H:%M")} | VIDYA AI')
    meta.runs[0].font.color.rgb = RGBColor(0x71, 0x8E, 0xA4)
    meta.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # Blank spacer line

    # --- Content ---
    # Split into lines and add each as a paragraph
    for line in content.split('\n'):
        if line.strip().startswith(('Q', 'A')) and ':' in line[:3]:
            # Q&A format: bold the Q/A label
            p = doc.add_paragraph()
            parts = line.split(':', 1)
            p.add_run(parts[0] + ':').bold = True
            if len(parts) > 1:
                p.add_run(parts[1])
        elif line.strip().startswith(('-', '*', chr(8226))):
            # Bullet point
            doc.add_paragraph(line.strip().lstrip('-* '), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())

    # --- Footer note ---
    doc.add_paragraph()
    doc.add_paragraph('-- Generated offline by VIDYA AI. Based on uploaded school materials. --')

    # Save to exports directory
    exports_dir = os.getenv('EXPORTS_DIR', './exports')
    os.makedirs(exports_dir, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_topic = topic.replace(' ', '_').replace('/', '-')[:30]
    filename = f'{grade}_{subject}_{safe_topic}_{note_type}_{timestamp}.docx'
    filepath = os.path.join(exports_dir, filename)

    doc.save(filepath)
    logger.info(f'Notes exported: {filepath}')
    return filepath